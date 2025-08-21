# api/models.py

from django.db import models
from django.contrib.auth.models import AbstractUser
import fitz  # PyMuPDF
import docx

# No changes to your custom User model
class User(AbstractUser):
    class Role(models.TextChoices):
        ADMIN = "ADMIN", "Admin"
        VIEWER = "VIEWER", "Viewer"
    role = models.CharField(max_length=50, choices=Role.choices, default=Role.VIEWER)

# NEW: UserProfile to manage admin status
class UserProfile(models.Model):
    user = models.OneToOneField(User, on_delete=models.CASCADE, related_name='profile')
    admin_approved = models.BooleanField(default=False)

    def __str__(self):
        return f"{self.user.username}'s Profile"

# NEW: Category model for document sorting
class Category(models.Model):
    name = models.CharField(max_length=100, unique=True)

    def __str__(self):
        return self.name

# UPDATED: Document model with new fields
class Document(models.Model):
    title = models.CharField(max_length=255)
    # NEW FIELDS
    description = models.TextField()
    cover_image = models.ImageField(upload_to='document_covers/')
    category = models.ForeignKey(Category, related_name='documents', on_delete=models.SET_NULL, null=True, blank=True)
    # EXISTING FIELDS
    file = models.FileField(upload_to='documents/')
    owner = models.ForeignKey(User, on_delete=models.CASCADE, related_name="owned_documents")
    uploaded_at = models.DateTimeField(auto_now_add=True)
    allowed_viewers = models.ManyToManyField(User, related_name='viewable_documents', blank=True)
    extracted_text = models.TextField(blank=True, null=True)
    watermark_enabled = models.BooleanField(default=True)

    def __str__(self):
        return self.title

    # No changes needed to your existing save method for text extraction
    def save(self, *args, **kwargs):
        is_new = self._state.adding
        super().save(*args, **kwargs)
        if is_new and self.file and not self.extracted_text:
            file_path = self.file.path
            text = ""
            try:
                if file_path.lower().endswith('.pdf'):
                    with fitz.open(file_path) as doc:
                        for page in doc:
                            text += page.get_text() + "\n"
                elif file_path.lower().endswith('.txt'):
                    with open(file_path, 'r', encoding='utf-8') as f:
                        text = f.read()
                elif file_path.lower().endswith('.docx'):
                    doc = docx.Document(file_path)
                    full_text = [para.text for para in doc.paragraphs]
                    text = '\n'.join(full_text)

                self.extracted_text = text
                super(Document, self).save(update_fields=['extracted_text'])
            except Exception as e:
                print(f"Error during text extraction for document ID {self.id}: {e}")

# NEW: DocumentRequest model
class DocumentRequest(models.Model):
    class AccessDuration(models.TextChoices):
        ONE_DAY = '1D', 'One Day'
        ONE_WEEK = '1W', 'One Week'
        FOREVER = 'F', 'Forever'

    class RequestStatus(models.TextChoices):
        PENDING = "PENDING", "Pending"
        APPROVED = "APPROVED", "Approved"
        DENIED = "DENIED", "Denied"

    document = models.ForeignKey(Document, related_name='requests', on_delete=models.CASCADE)
    requester = models.ForeignKey(User, related_name='requests_made', on_delete=models.CASCADE)
    duration = models.CharField(max_length=2, choices=AccessDuration.choices, default=AccessDuration.FOREVER)
    status = models.CharField(max_length=20, choices=RequestStatus.choices, default=RequestStatus.PENDING)
    requested_at = models.DateTimeField(auto_now_add=True)

    def __str__(self):
        return f"Request for '{self.document.title}' by {self.requester.username}"